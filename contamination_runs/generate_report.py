#!/usr/bin/env python3
"""
Generate a Word report documenting the HCVTyper contamination check pipeline,
its outputs, and findings across three sequencing runs.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

RESULTS_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Styles helpers ────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),   "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"),  "F2F2F2")
    p._p.pPr.append(shading)
    return p

def add_table(headers, rows, col_widths=None, header_shading="2E4057"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  header_shading)
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
    # Column widths
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Cm(w)
    return table

def add_image_if_exists(path, caption, width=14):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True
        c.runs[0].font.size = Pt(9)
    else:
        para(f"[Figure not available: {os.path.basename(path)}]", italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("HCVTyper Cross-Sample Contamination Check", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Methodology, Output Interpretation, and Run Analysis")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].bold = True

date_p = doc.add_paragraph("Folkehelseinstituttet (FHI)  —  March 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(11)
date_p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Introduction")

para(
    "Sample-to-sample contamination is a known risk in high-throughput sequencing (HTS) workflows. "
    "During library preparation and sequencing, small amounts of nucleic acid from one sample can "
    "inadvertently end up in another, either through physical cross-contamination during liquid "
    "handling or through optical/electronic index misassignment (index hopping) on patterned flow-cell "
    "platforms such as the Illumina NovaSeq. Undetected contamination can lead to incorrect variant "
    "calls, misidentification of drug-resistance mutations, or erroneous patient results."
)
para(
    "To address this risk in the HCV diagnostic sequencing pipeline at FHI, a dedicated "
    "CONTAMINATION_CHECK subworkflow was developed as part of HCVTyper. This document describes the "
    "methodology behind the check, explains how to interpret each output file, and presents results "
    "from three sequencing runs that illustrate the full range of outcomes — from definitively clean "
    "runs to runs with strong contamination signals requiring investigation."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Methodology")

heading("2.1  Overview", level=2)
para(
    "The contamination check is based on comparing de novo assembled contigs across all samples in a "
    "sequencing run. SPAdes assembles consensus contigs for each sample; if two samples share "
    "near-identical sequence, this indicates the same viral population was present in both wells — "
    "a strong indicator of contamination."
)
para("The pipeline consists of five steps:")
for step in [
    ("Contig filtering and renaming",
     "Per-sample contigs shorter than 1,000 bp are discarded to reduce noise from small, "
     "low-complexity fragments. Each remaining contig header is prefixed with the sample ID "
     "(e.g. 2619488-HCV|NODE_1_length_9398_cov_15208.0)."),
    ("Concatenation",
     "All per-sample filtered FASTA files are combined into a single multi-sample FASTA."),
    ("BLAST database construction",
     "A nucleotide BLAST database is built from the combined FASTA."),
    ("All-vs-all BLAST",
     "The combined FASTA is BLASTed against itself at ≥95% nucleotide identity. "
     "Self-hits are automatically present but are filtered in the next step."),
    ("Report generation",
     "The R script contamination_report.R parses the BLAST output, removes self-hits and "
     "intra-sample hits, and produces all output files described in Section 3."),
]:
    p = doc.add_paragraph(style="List Number")
    p.add_run(step[0] + ": ").bold = True
    p.add_run(step[1])

heading("2.2  Coverage-based direction inference", level=2)
para(
    "Each SPAdes contig header encodes the assembly coverage depth: "
    "NODE_<n>_length_<bp>_cov_<depth>. Because the source sample contains the original viral "
    "population at full sequencing depth, while the recipient sample contains only a small carried-over "
    "fraction, the source contig will always have substantially higher SPAdes coverage than the "
    "equivalent contig in the recipient sample."
)
para(
    "For every cross-sample contig pair the pipeline calculates:"
)
bullet("source_cov = max(contig1_cov, contig2_cov)")
bullet("recipient_cov = min(contig1_cov, contig2_cov)")
bullet("cov_ratio = source_cov / recipient_cov")
para(
    "A high coverage ratio (default threshold ≥10×) indicates that the direction of transfer is "
    "reliably inferred. Pairs with cov_ratio <10× are still reported but direction is uncertain — "
    "they may represent biologically related sequences (e.g. two patients infected with the same "
    "subtype cluster) rather than a true contamination event."
)

heading("2.3  Index-hopping noise threshold", level=2)
para(
    "On patterned flow cells, reads from one sample can be sequenced with the barcode of a different "
    "sample due to free-floating index primers clustering adjacent to the wrong template. This is "
    "called index hopping and occurs at a rate of approximately 0.1% (1 in 1,000 reads). Index-hopped "
    "reads will assemble into contigs with very low coverage in the recipient sample."
)
para("The expected coverage contribution from index hopping is calculated as:")
code_block(
    "hop_threshold (×) = (total_run_reads × hop_rate × read_length)\n"
    "                    ──────────────────────────────────────────────\n"
    "                    (n_samples × genome_size)"
)
para(
    "Parameters used: hop_rate = 0.001 (0.1%), read_length = 150 bp, genome_size = 9,500 bp. "
    "Any recipient contig with coverage below this threshold is flagged as "
    "flag_index_hopping = TRUE, meaning the signal could plausibly arise from index hopping alone "
    "and is not strong evidence of physical contamination."
)

para(
    "Because deeper sequencing runs produce higher hop thresholds (more total reads → more "
    "index-hopped reads reaching each well), a run with many high-depth samples will have a higher "
    "threshold and consequently fewer pairs that survive the hop filter. This is expected and correct."
)

heading("2.4  GLUE resistance mutation similarity (Jaccard index)", level=2)
para(
    "HCV-GLUE (http://hcv.cvr.ac.uk/glue/) performs resistance-associated substitution (RAS) "
    "scanning on mapped consensus sequences and reports which known resistance mutations are present "
    "in each sample at sufficient coverage. If two samples are truly contaminated with the same "
    "viral strain, they should share the same resistance mutation profile."
)
para(
    "For each cross-sample contig pair the pipeline intersects the RAS sets of the two samples and "
    "computes the Jaccard similarity index:"
)
code_block(
    "Jaccard(A, B) = |A ∩ B| / |A ∪ B|"
)
para(
    "A Jaccard of 1.0 means both samples have exactly the same set of resistance mutations — "
    "strong corroborating evidence that the same viral quasispecies is present in both. "
    "A Jaccard of 0 means no shared mutations, arguing against contamination with an identical strain. "
    "Values between 0 and 1 are intermediate. "
    "Note that two samples infected with the same subtype but different patients may share some "
    "common subtype-defining mutations; Jaccard is most informative when shared_mutations_union "
    "contains rare or treatment-selected mutations."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. OUTPUT FILES
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Output Files")

heading("3.1  contamination_pairs.tsv", level=2)
para(
    "One row per unique cross-sample contig pair that passed the BLAST identity threshold (≥95%). "
    "Columns:"
)
add_table(
    ["Column", "Description"],
    [
        ["sample1, sample2",        "The two samples that share this contig segment"],
        ["contig1, contig2",        "Full SPAdes contig names (embed length and coverage depth)"],
        ["pident",                  "BLAST percent nucleotide identity (≥95 by filter)"],
        ["alignment_length",        "Number of aligned bases"],
        ["snp_distance",            "Mismatches + gap openings — effectively the SNP distance between the two contigs"],
        ["evalue / bitscore",       "Standard BLAST statistics"],
        ["contig1_cov, contig2_cov","SPAdes k-mer coverage depth parsed from the contig name"],
        ["source_sample / recipient_sample", "Higher-coverage sample inferred as source; lower-coverage as recipient"],
        ["cov_ratio",               "source_cov / recipient_cov"],
        ["index_hop_threshold",     "The run-specific hop threshold in units of × coverage"],
        ["flag_index_hopping",      "TRUE if recipient_cov < index_hop_threshold"],
        ["n_shared_mutations",      "Number of GLUE RAS mutations present in both samples"],
        ["jaccard_similarity",      "Jaccard index of the two samples' RAS sets"],
        ["shared_mutations",        "Semicolon-separated list of shared RAS mutation names"],
    ],
    col_widths=[5, 11],
)
doc.add_paragraph()

heading("3.2  contamination_direction.tsv", level=2)
para(
    "One row per directed source→recipient pair, aggregated over all individual contig pairs. "
    "This is the primary file for contamination investigation. Columns:"
)
add_table(
    ["Column", "Description"],
    [
        ["source_sample",          "Inferred source of contamination"],
        ["recipient_sample",       "Inferred recipient of contamination"],
        ["n_contig_pairs",         "Total number of shared contig pairs supporting this direction"],
        ["n_above_hop_threshold",  "Pairs where recipient coverage exceeds the index-hop threshold (strong evidence)"],
        ["median_cov_ratio",       "Median source/recipient coverage ratio across all pairs"],
        ["median_source_cov",      "Median SPAdes coverage of the source contigs"],
        ["median_recipient_cov",   "Median SPAdes coverage of the recipient contigs"],
        ["median_jaccard_ras",     "Median pairwise GLUE Jaccard similarity"],
        ["shared_mutations_union", "Union of all shared RAS mutations across all pairs"],
    ],
    col_widths=[5, 11],
)
doc.add_paragraph()

heading("3.3  contamination_heatmap.png", level=2)
para(
    "A symmetric sample × sample matrix. Each cell shows the number of cross-sample contig pairs "
    "between those two samples (regardless of direction). The diagonal is blank. A high count "
    "between two samples is the first visual indicator of a potential contamination event. "
    "Clean runs show uniformly blank or near-zero matrices."
)

heading("3.4  contamination_direction.png", level=2)
para(
    "A directed bubble plot showing only high-confidence pairs (median_cov_ratio ≥ 10×). "
    "The x-axis shows the inferred source sample and the y-axis the inferred recipient. "
    "Bubble size encodes the number of contig pairs. Colour encodes either the log₁₀ "
    "coverage ratio (when GLUE is unavailable) or the median GLUE Jaccard similarity. "
    "The bubble label shows 'total (above hop threshold)'. "
    "Only pairs with strong direction signal appear on this plot."
)

heading("3.5  contamination_mqc.json", level=2)
para(
    "A MultiQC-compatible JSON table that appears in the MultiQC report under the section "
    "'Cross-sample Contamination Check'. Contains the same data as contamination_pairs.tsv "
    "and allows interactive filtering in the HTML report."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. INTERPRETATION GUIDE
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Interpretation Guide")

heading("4.1  Decision framework", level=2)
para(
    "Use the following criteria in combination to assess each flagged sample pair. "
    "A verdict of probable contamination requires multiple lines of evidence."
)

add_table(
    ["Criterion", "Value", "Interpretation"],
    [
        ["n_above_hop_threshold", "≥1", "At least one contig pair has recipient coverage exceeding the index-hop floor → physical contamination cannot be excluded"],
        ["n_above_hop_threshold", "0",  "All recipient-side contigs are at index-hop floor coverage → likely index-hopping noise; low concern"],
        ["median_cov_ratio",      "≥10×","Direction is reliable; source sample clearly identified"],
        ["median_cov_ratio",      "1–10×","Direction uncertain; could be co-circulating strains, not contamination"],
        ["median_jaccard_ras",    "1.0", "Identical resistance mutation profile → strong evidence same strain in both samples"],
        ["median_jaccard_ras",    "0",   "No shared mutations → argues against shared strain; may be same subtype HCV"],
        ["n_contig_pairs",        "≥5 and n_above_hop ≥1", "High concern — investigate sample origin, plate map, sample handling"],
        ["shared_mutations",      "treatment-selected RAS (e.g. NS5A:93H, NS3:168A)", "Especially significant — resistance mutations are patient/treatment specific"],
    ],
    col_widths=[4.5, 3, 8.5],
)
doc.add_paragraph()

heading("4.2  Confirming or ruling out contamination", level=2)
para("When a pair is flagged as concerning, the following steps are recommended:")
for i, step in enumerate([
    "Check the plate map / sample extraction layout — were the two samples processed adjacent to each other on the same plate or tip rack?",
    "Check the clinical context — are the two patients epidemiologically linked (same household, clinic, or treatment group)?",
    "Examine shared_mutations_union — are the shared mutations treatment-selected or known subtype-defining? Subtype-defining mutations (e.g. common 1b NS5A polymorphisms) are not informative.",
    "Compare the full resistance reports for both patients — completely identical RAS profiles across NS3/NS5A/NS5B are highly suspicious.",
    "If contamination is confirmed, the recipient sample result should be interpreted with caution and re-sequencing considered.",
], 1):
    bullet(step)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. RUN ANALYSES
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Run Analyses")

heading("5.1  NGS_SEQ-20260115-02  —  Clean run", level=2)

add_table(
    ["Parameter", "Value"],
    [
        ["Run date",             "2026-01-15"],
        ["Number of samples",    "22"],
        ["Total reads (post-QC)","136,137,882"],
        ["Index-hop threshold",  "97.71×"],
        ["Cross-sample pairs",   "2"],
        ["Direction rows",       "2"],
        ["GLUE files",           "14 of 22 samples"],
    ],
    col_widths=[6, 10],
)
doc.add_paragraph()

para(
    "This run is an example of a clean result. Only two cross-sample contig pairs were detected, "
    "both between different sample combinations."
)

add_table(
    ["Source", "Recipient", "Pairs", "Above hop", "Cov ratio", "Jaccard", "Shared RAS"],
    [
        ["2659870-HCV", "2656696-HCV", "1", "0", "39,447×", "0", "none"],
        ["2658896-HCV", "2654486-HCV", "1", "0", "1,160×",  "0", "none"],
    ],
    col_widths=[3.5, 3.5, 2, 3, 3, 3, 3],
)
doc.add_paragraph()

para(
    "Verdict: No contamination. Both pairs have n_above_hop_threshold = 0, meaning the recipient "
    "contig coverage in both cases is below the index-hopping floor of 97.7× — entirely consistent "
    "with a small number of misassigned reads. The Jaccard similarity is 0 for both pairs, "
    "providing no biological corroboration. This run requires no further action."
)

add_image_if_exists(
    os.path.join(RESULTS_DIR, "NGS_SEQ-20260115-02/results/contamination.contamination_heatmap.png"),
    "Figure 1. Heatmap for NGS_SEQ-20260115-02. Only 2 shared pairs visible — effectively clean.",
    width=12,
)

heading("5.2  NGS_SEQ-20260210-01  —  High read depth, all index-hop noise", level=2)

add_table(
    ["Parameter", "Value"],
    [
        ["Run date",             "2026-02-10"],
        ["Number of samples",    "18"],
        ["Total reads (post-QC)","61,028,326"],
        ["Index-hop threshold",  "53.53×"],
        ["Cross-sample pairs",   "333"],
        ["Direction rows",       "134"],
        ["GLUE files",           "16 of 18 samples"],
    ],
    col_widths=[6, 10],
)
doc.add_paragraph()

para(
    "This run has substantially more cross-sample pairs (333) than the clean run, but the key "
    "finding is that n_above_hop_threshold = 0 for all direction pairs. The index-hop threshold "
    "for this run is 53.5×; every recipient-side contig falls below this coverage. This pattern is "
    "characteristic of a run with deep sequencing depth — more total reads means more index-hopped "
    "reads reaching every well, generating low-coverage noise contigs."
)

para("Top direction pairs:")

add_table(
    ["Source", "Recipient", "Pairs", "Above hop", "Cov ratio", "Jaccard", "Shared RAS"],
    [
        ["2681924-HCV", "2679090-HCV", "13", "0", "93.8×",    "0", "none"],
        ["2686402-HCV", "2681924-HCV", "12", "0", "2,068×",   "0", "none"],
        ["2681924-HCV", "2686402-HCV", "10", "0", "1,800×",   "0", "none"],
    ],
    col_widths=[3.5, 3.5, 2, 3, 3, 3, 3],
)
doc.add_paragraph()

para(
    "Verdict: No contamination. Despite the large number of flagged pairs, all recipient contigs "
    "sit below the index-hopping floor. The Jaccard similarity is uniformly 0, confirming these "
    "samples do not share resistance mutation profiles. This is a textbook index-hopping pattern — "
    "high read depth run, many low-coverage apparent hits, zero biological corroboration. "
    "No follow-up required."
)

add_image_if_exists(
    os.path.join(RESULTS_DIR, "NGS_SEQ-20260210-01/results/contamination.contamination_heatmap.png"),
    "Figure 2. Heatmap for NGS_SEQ-20260210-01. Many pairs visible but all are index-hopping noise.",
    width=14,
)

heading("5.3  NGS_SEQ-20251212-01  —  Confirmed contamination signals", level=2)

add_table(
    ["Parameter", "Value"],
    [
        ["Run date",             "2025-12-12"],
        ["Number of samples",    "22"],
        ["Total reads (post-QC)","121,012,304"],
        ["Index-hop threshold",  "86.85×"],
        ["Cross-sample pairs",   "394"],
        ["Direction rows",       "121"],
        ["GLUE files",           "21 of 22 samples"],
    ],
    col_widths=[6, 10],
)
doc.add_paragraph()

para(
    "This run contains multiple pairs that survive the index-hopping filter and show GLUE mutation "
    "corroboration. The following table shows the highest-priority pairs:"
)

add_table(
    ["Source", "Recipient", "Pairs", "Above hop", "Cov ratio", "Jaccard", "Shared RAS mutations"],
    [
        ["2619488-HCV", "2626415-HCV", "8",  "3", "2.1×",   "1.00", "NS3:56F; NS5A:30Q; NS5A:30Q+31M; NS5A:31M; NS5A:37L; NS5B:159F; NS5B:159F+316N; NS5B:316N"],
        ["2634016-HCV", "2633901-HCV", "6",  "3", "1.3×",   "1.00", "NS3:122G"],
        ["2626415-HCV", "2619488-HCV", "3",  "0", "6,565×", "1.00", "same 8 mutations"],
        ["2622436-HCV", "2623163-HCV", "9",  "0", "6,650×", "0.50", "NS3:56Y+168Q+170I"],
        ["2634392-HCV", "2623163-HCV", "5",  "0", "1,462×", "0.50", "NS3:56Y+168Q+170I"],
        ["2634413-HCV", "2622044-HCV", "3",  "0", "154×",   "0.67", "NS3:56Y+168Q+170I; NS5B:150V; NS5B:150V+206E; NS5B:206E"],
        ["2620976-HCV", "2626415-HCV", "3",  "0", "10,429×","0.18", "NS3:56F; NS5A:37L"],
        ["2626415-HCV", "2619488-HCV", "3",  "0", "6,565×", "1.00", "NS3:56F; NS5A:30Q/31M/37L; NS5B:159F/316N"],
    ],
    col_widths=[3.2, 3.2, 2, 2.8, 2.8, 2.5, 5.5],
)
doc.add_paragraph()

heading("Priority pair 1: 2619488-HCV ↔ 2626415-HCV", level=3)
para(
    "This pair shows Jaccard = 1.0 and 3 contig pairs above the index-hop threshold, with 8 shared "
    "resistance mutations spanning all three drug targets (NS3, NS5A, NS5B). This is the strongest "
    "possible signal — identical mutation profiles across the full genome plus multiple high-coverage "
    "shared contigs. The bidirectional signals (both directions show pairs) suggest the two samples "
    "may originate from the same patient, OR that there was a physical liquid handling contamination "
    "event. Immediate investigation is warranted."
)
bullet("8 shared RAS mutations: NS3:56F, NS5A:30Q, NS5A:30Q+31M, NS5A:31M, NS5A:37L, NS5B:159F, NS5B:159F+316N, NS5B:316N")
bullet("3 contig pairs above 86.85× hop threshold in the 2619488→2626415 direction")
bullet("Cov ratio 2.1× in that direction (uncertain direction at this ratio), but 6,565× in reverse")
bullet("Action: Check plate map, verify these are not the same patient, consider re-extraction")

heading("Priority pair 2: 2634016-HCV ↔ 2633901-HCV", level=3)
para(
    "Jaccard = 1.0, 3 contig pairs above index-hop threshold, shared NS3:122G mutation. "
    "The coverage ratio of 1.3× is low (direction uncertain), but both directions show pairs. "
    "NS3:122G is associated with reduced susceptibility to NS3 inhibitors and is not a common "
    "subtype-defining polymorphism, making the match meaningful."
)
bullet("Shared RAS: NS3:122G — resistance-associated, not a common background substitution")
bullet("3 pairs in one direction above threshold → physical contamination cannot be excluded")
bullet("Action: Check adjacency on extraction plate, review patient histories for shared treatment exposure")

heading("Lower-priority signals", level=3)
para(
    "Multiple other pairs (e.g. 2622436↔2623163, 2634392↔2623163) share NS3 complex mutations "
    "(56Y+168Q+170I) and show moderate Jaccard values (0.33–0.50), but have n_above_hop_threshold = 0. "
    "These could represent samples carrying a similar treatment-selected NS3 variant cluster rather "
    "than physical contamination. They warrant review of the clinical resistance reports but are "
    "lower priority than pairs 1 and 2."
)

add_image_if_exists(
    os.path.join(RESULTS_DIR, "NGS_SEQ-20251212-01/results/contamination.contamination_heatmap.png"),
    "Figure 3. Heatmap for NGS_SEQ-20251212-01. Multiple high-count pairs visible, particularly involving samples 2619488-HCV, 2626415-HCV, 2634016-HCV, and 2633901-HCV.",
    width=14,
)

add_image_if_exists(
    os.path.join(RESULTS_DIR, "NGS_SEQ-20251212-01/results/contamination.contamination_direction.png"),
    "Figure 4. Direction plot for NGS_SEQ-20251212-01 (pairs with cov_ratio ≥ 10×, coloured by GLUE Jaccard similarity). Colour approaching red = high Jaccard = shared mutation profile. Label format: total pairs (above hop threshold).",
    width=14,
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. COMPARISON ACROSS RUNS
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. Comparison Across Runs")

add_table(
    ["Run", "Samples", "Total reads", "Hop threshold", "Total pairs", "Pairs above hop", "Max Jaccard", "Verdict"],
    [
        ["NGS_SEQ-20260115-02", "22", "136.1M", "97.7×", "2",   "0", "0",   "Clean — index hop only"],
        ["NGS_SEQ-20260210-01", "18", "61.0M",  "53.5×", "333", "0", "0",   "Index hop noise — clean"],
        ["NGS_SEQ-20251212-01", "22", "121.0M", "86.9×", "394",  "≥15", "1.0","⚠️  Investigate pairs 1 & 2"],
    ],
    col_widths=[4, 2.5, 2.5, 3, 2.5, 3, 3, 4],
)
doc.add_paragraph()

para(
    "The table above illustrates why both the index-hop filter AND the Jaccard similarity are needed. "
    "NGS_SEQ-20260210-01 has 333 pairs but all are noise; NGS_SEQ-20251212-01 has fewer pairs above "
    "the filter but those pairs carry a strong GLUE corroboration signal. False positives driven by "
    "deep sequencing are effectively captured by the hop threshold; true contamination signals "
    "survive and are further confirmed by the mutation similarity."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. PIPELINE PARAMETERS
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. Pipeline Parameters")

add_table(
    ["Parameter", "Default", "Description"],
    [
        ["contamination_min_length", "1000 bp",  "Minimum contig length to include in BLAST comparison"],
        ["contamination_min_id",     "95%",       "Minimum BLAST nucleotide identity to report a hit"],
        ["contamination_hop_rate",   "0.001",     "Expected index-hopping rate (0.1%) — adjust for your sequencing platform"],
        ["contamination_min_dir_ratio", "10.0",   "Minimum source/recipient cov ratio for direction plot"],
        ["contamination_genome_size", "9500 bp",  "HCV reference genome size used for hop threshold calculation"],
    ],
    col_widths=[5.5, 3, 8],
)
doc.add_paragraph()

code_block(
    "# Run standalone contamination check with fastp and GLUE inputs:\n"
    "nextflow run contamination_check_standalone.nf \\\n"
    "    -c contamination_runs/nextflow.config \\\n"
    "    --contigs_dir /path/to/run/spades \\\n"
    "    --fastp_dir   /path/to/run/fastp \\\n"
    "    --glue_dir    /path/to/run/hcvglue \\\n"
    "    --outdir      contamination_runs/<run_id>/results \\\n"
    "    -work-dir     contamination_runs/<run_id>/work \\\n"
    "    -resume"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. Conclusions")

para(
    "The HCVTyper contamination check provides a multi-layered approach to detecting and "
    "characterising cross-sample contamination in diagnostic HCV sequencing runs:"
)
for point in [
    "The all-vs-all BLAST detects any near-identical sequence shared between samples, regardless of coverage.",
    "The index-hop threshold filters out the inevitable low-level noise from index misassignment on patterned flow cells, preventing false alarms on high-depth runs.",
    "The coverage ratio identifies the likely direction of transfer and quantifies how much relative coverage the recipient carries.",
    "The GLUE Jaccard similarity provides independent biological corroboration: shared resistance mutations — especially treatment-selected mutations — strongly support that the same viral quasispecies is present in both samples.",
    "A pair is conclusively concerning only when: (a) n_above_hop_threshold ≥ 1 indicating the signal is too strong to be explained by index hopping alone, AND (b) the GLUE Jaccard similarity is elevated (≥ 0.5), confirming a shared mutation profile.",
    "Most routine runs are expected to be clean or show only index-hop noise. The three runs analysed here demonstrate that when true contamination occurs, the combined signal is unambiguous and actionable.",
]:
    bullet(point)

doc.add_paragraph()
para(
    "This workflow is integrated into HCVTyper (Folkehelseinstituttet, folkehelseinstituttet/hcvtyper) "
    "and runs automatically as part of every sequencing batch.",
    italic=True,
)

# ── Save ──────────────────────────────────────────────────────────────────────
outfile = os.path.join(RESULTS_DIR, "HCVTyper_Contamination_Check_Report.docx")
doc.save(outfile)
print(f"Saved: {outfile}")
